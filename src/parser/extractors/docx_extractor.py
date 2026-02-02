"""DOCX text extraction. UTF-8, preserves Arabic diacritics (harakat)."""

from pathlib import Path
from typing import Any

from docx import Document


def extract(path: str | Path) -> dict[str, Any]:
    """
    Extract text from a DOCX file.
    Returns: {"raw_text": str, "pages_or_sections": list[dict]} with paragraph-level structure.
    Preserves UTF-8 and Arabic diacritics.
    """
    path = Path(path)
    if not path.exists():
        raise FileNotFoundError(f"DOCX not found: {path}")
    if path.suffix.lower() not in (".docx", ".doc"):
        raise ValueError(f"Expected .docx/.doc file, got {path.suffix}")

    doc = Document(str(path))
    pages_or_sections: list[dict[str, Any]] = []
    raw_parts: list[str] = []

    for i, para in enumerate(doc.paragraphs):
        text = para.text or ""
        raw_parts.append(text)
        pages_or_sections.append({"index": i, "text": text, "style": para.style.name if para.style else None})

    # Optional: extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " ".join(cell.text or "" for cell in row.cells)
            if row_text.strip():
                raw_parts.append(row_text)
                pages_or_sections.append({"type": "table_cell", "text": row_text})

    raw_text = "\n\n".join(raw_parts)
    return {"raw_text": raw_text, "pages_or_sections": pages_or_sections}
